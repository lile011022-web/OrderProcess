from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "球星卡采购接单对账系统_PRD_v0.2.docx"


def set_run(run, size=11, bold=False, color="000000"):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def qn(tag):
    from docx.oxml.ns import qn as _qn

    return _qn(tag)


def add_paragraph(doc, text="", style=None, size=11, bold=False, color="000000", after=6):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    if text:
      run = p.add_run(text)
      set_run(run, size=size, bold=bold, color=color)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18 if level == 1 else 14)
    p.paragraph_format.space_after = Pt(8 if level == 1 else 6)
    run = p.add_run(text)
    set_run(run, size=16 if level == 1 else 13, bold=True, color="2E74B5" if level == 1 else "1F4D78")
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        run = p.add_run(item)
        set_run(run)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        run = p.add_run(item)
        set_run(run)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    from docx.oxml import OxmlElement

    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    hdr = table.rows[0].cells
    for idx, text in enumerate(headers):
        hdr[idx].width = Inches(widths[idx])
        hdr[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        shade_cell(hdr[idx], "E8EEF5")
        p = hdr[idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        set_run(run, bold=True, color="0B2545")
    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            cells[idx].width = Inches(widths[idx])
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(str(text))
            set_run(run, size=10)
    add_paragraph(doc, "")
    return table


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = Pt(11)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run("球星卡采购接单对账系统 PRD")
    set_run(run, size=24, bold=True, color="0B2545")
    add_paragraph(doc, "版本：v0.2    更新时间：2026-07-02    阶段：前端原型", size=10, color="555555", after=12)

    add_heading(doc, "1. 产品目标")
    add_paragraph(doc, "本系统用于管理球星卡采购任务、买手接单采购、采购回填、包裹物流、仓库收货、付款确认、异常处理、基础资料和对账统计。第一阶段只实现可交互前端原型，不连接真实后端，不调用真实物流接口，统一使用 mock data 模拟业务。")

    add_heading(doc, "2. 本次修订结论")
    add_bullets(doc, [
        "移除单独的“费用规则”设置入口，保留必要的集中计算能力供成本页面复用。",
        "运单号由系统自动识别 UPS、FedEx、USPS，买手不再手动选择快递公司。",
        "预计到达时间改由仓库端在待确认包裹中填写或更新。",
        "收货仓库由管理员维护，买手通过下拉框选择已启用仓库；收货人由买手填写姓名。",
        "管理员付款、买手收款均按采购数量处理，支持全选或按部分数量处理。",
        "仓库照片入口合并到已收货包裹和超时核查的对应单号下。",
        "商品库只管理商品基础资料；仓库地址采用美国地址字段维护。",
        "客户页面和管理员页面分隔，客户只维护自己的商品资料和仓库地址提交。",
        "纯图标按钮提供悬停说明，当前选中菜单与 hover 状态明显区分。",
    ])

    add_heading(doc, "3. 角色与权限")
    add_table(doc, ["角色", "可见范围", "主要动作"], [
        ["管理员", "采购、买手、包裹、财务、报表、基础资料", "发布任务、审核、付款、维护基础资料"],
        ["买手", "任务大厅、我的接单、采购回填、收款状态", "接单、回填采购和物流信息、查看收款"],
        ["仓库", "待确认包裹、已收货包裹、超时待核查、仓库操作费", "填写 ETA、确认收货、上传照片、核查超时"],
        ["客户", "我的商品资料、我的仓库地址、提交审核状态", "提交商品、提交美国地址、查看审核结果"],
    ], [1.0, 2.7, 2.8])

    add_heading(doc, "4. 核心业务流程")
    add_numbered(doc, [
        "管理员发布采购任务，填写商品、数量、目标价、采购要求和收货信息。",
        "买手在任务大厅接单并完成采购。",
        "买手回填实际采购数量、单价、费用、平台订单号、运单号、收货仓库和收货人。",
        "系统根据运单号识别快递公司，并提供官网查询跳转。",
        "管理员审核买手回填记录。",
        "财务按采购数量付款；付款后金额进入“已付待确认”。",
        "仓库在待确认包裹中填写或更新预计送达时间，并按运单号确认收货。",
        "仓库确认实际收到后，相关金额才转为实际入库成本。",
        "已收货包裹照片在对应运单号详情中上传和查看。",
        "超时未收货包裹进入核查流程，仓库可上传官网截图、包裹照片或异常照片。",
    ])

    add_heading(doc, "5. 页面需求")
    add_table(doc, ["页面", "关键需求"], [
        ["买手采购回填", "买手填写采购数量、价格、费用、平台订单、运单号、收货仓库和收货人；系统识别快递；不填写预计到达时间。"],
        ["买手收款 / 管理员付款", "买手端展示为收款，管理员端展示为付款；支持按采购数量全选或部分处理；付款后进入已付待确认。"],
        ["仓库待确认包裹", "仓库填写预计送达时间；扫描或输入运单号；确认实收数量并上传收货照片。"],
        ["已收货包裹", "专用业务页面；每个运单号提供上传和查看外箱、商品、异常补充照片入口。"],
        ["超时待核查", "支持继续等待、已收到、未收到、物流异常；上传官网截图和异常反馈照片。"],
        ["商品库", "只维护商品编码、名称、分类、品牌、规格、资料来源、参考价和状态。"],
        ["仓库地址", "管理员和客户提交维护；使用 Address Line、City、State、ZIP Code、Country 等美国地址字段。"],
        ["客户资料", "客户登录后只看到自己的商品资料和仓库地址；不展示管理员审核、启用、停用等全局管理动作。"],
    ], [1.55, 4.95])

    add_heading(doc, "6. 集中规则")
    add_bullets(doc, [
        "运单识别：集中维护 UPS、FedEx、USPS 识别规则和官网链接，未知单号展示“待识别”。",
        "付款与入库成本：付款只能进入“已付待确认”，仓库确认收到后才进入实际入库成本。",
        "异常处理：记录原因、责任方、金额、处理方式、状态、备注和凭证，不混入正常入库成本。",
    ])

    add_heading(doc, "7. 原型验收标准")
    add_bullets(doc, [
        "未登录访问后台跳转到 /login，四种角色可模拟登录且菜单不同。",
        "客户使用 customer / 123456 登录后只能看到客户资料菜单。",
        "客户页面与管理员基础资料页面分隔，客户不看到全量商品库或全量仓库地址。",
        "买手菜单不会出现同一路径导致两个入口同时高亮。",
        "买手回填只填运单号，系统显示快递识别结果。",
        "买手收款/管理员付款支持按采购数量全选或部分处理。",
        "仓库待确认中可以填写预计送达时间。",
        "已收货包裹中可以按对应单号上传照片。",
        "超时待核查中可以上传官网截图和异常反馈照片。",
        "商品库和仓库地址页面展示真实业务字段。",
        "纯图标按钮具备悬停说明。",
    ])

    add_heading(doc, "8. 后续待办")
    add_bullets(doc, [
        "接入真实后端后，将 mock data 替换为 API 数据访问层。",
        "自动读取官网预计到达时间需评估物流官网接口、反爬限制、账号权限和数据授权。",
        "付款、收货、异常处理后续需要增加真实状态流转和审计日志。",
    ])

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
